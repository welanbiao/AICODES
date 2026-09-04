# -*- coding: utf-8 -*-
"""Export AIkp filing markdown to Word (.docx) with embedded screenshots."""
from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
import markdown as md

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "防侵权备案_产品设定与开发说明.md"
OUT_PATH = ROOT / "docs" / "防侵权备案_产品设定与开发说明.docx"
IMG_BASE = ROOT / "docs"


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text_with_inline(paragraph, node, base_size=11, bold=False):
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, size=base_size, bold=bold)
        return
    if not isinstance(node, Tag):
        return
    name = node.name.lower()
    if name in ("strong", "b"):
        for child in node.children:
            add_text_with_inline(paragraph, child, base_size, True)
    elif name in ("em", "i"):
        for child in node.children:
            add_text_with_inline(paragraph, child, base_size, bold)
            if paragraph.runs:
                paragraph.runs[-1].italic = True
    elif name == "code":
        run = paragraph.add_run(node.get_text())
        set_run_font(run, name="Consolas", size=9, bold=bold)
    elif name == "a":
        for child in node.children:
            add_text_with_inline(paragraph, child, base_size, bold)
    elif name == "br":
        paragraph.add_run().add_break()
    else:
        for child in node.children:
            add_text_with_inline(paragraph, child, base_size, bold)


def add_paragraph_from_tag(doc, tag, style=None, size=11, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    for child in tag.children:
        add_text_with_inline(p, child, base_size=size)
    if not p.runs and tag.get_text(strip=True):
        run = p.add_run(tag.get_text(" ", strip=True))
        set_run_font(run, size=size)
    return p


def add_heading(doc, text, level):
    # Word heading levels 1-3
    level = max(1, min(level, 3))
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_run_font(run, size=sizes[level], bold=True, color=RGBColor(0x0B, 0x12, 0x20))
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    cols = max(len(r.find_all(["th", "td"])) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            if j < len(cells):
                txt = cells[j].get_text(" ", strip=True)
                is_header = cells[j].name == "th" or i == 0
                run = p.add_run(txt)
                set_run_font(run, size=9, bold=is_header)
    doc.add_paragraph()


def add_image(doc, src: str, alt: str = ""):
    # markdown images are like screenshots/01_auth.png relative to docs/
    path = (IMG_BASE / src).resolve()
    if not path.exists():
        # try raw src
        path = (ROOT / src).resolve()
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[缺少图片: {src}]")
        set_run_font(run, size=10, color=RGBColor(0xB2, 0x3A, 0x48))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # phone screenshots are tall; keep readable width
    run.add_picture(str(path), width=Cm(9.5))
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt)
        set_run_font(r, size=9, color=RGBColor(0x5A, 0x66, 0x78))


def add_code_block(doc, text: str):
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(line if line else " ")
        set_run_font(run, name="Consolas", size=8)
    doc.add_paragraph()


def walk(doc, soup: BeautifulSoup):
    body = soup.body if soup.body else soup
    for el in body.children:
        if isinstance(el, NavigableString):
            continue
        if not isinstance(el, Tag):
            continue
        name = el.name.lower()
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            add_heading(doc, el.get_text(" ", strip=True), level)
        elif name == "p":
            img = el.find("img")
            if img and not el.get_text(strip=True).replace(img.get("alt") or "", ""):
                add_image(doc, img.get("src") or "", img.get("alt") or "")
            else:
                # mixed p with possible img
                if img:
                    add_paragraph_from_tag(doc, el, size=11)
                    add_image(doc, img.get("src") or "", img.get("alt") or "")
                else:
                    add_paragraph_from_tag(doc, el, size=11)
        elif name == "ul":
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                for child in li.children:
                    add_text_with_inline(p, child, 10)
        elif name == "ol":
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(2)
                for child in li.children:
                    add_text_with_inline(p, child, 10)
        elif name == "table":
            add_table(doc, el)
        elif name == "pre":
            code = el.get_text()
            add_code_block(doc, code.rstrip("\n"))
        elif name == "blockquote":
            p = add_paragraph_from_tag(doc, el, size=10)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x5A, 0x66, 0x78)
        elif name == "hr":
            p = doc.add_paragraph("─" * 28)
            for run in p.runs:
                set_run_font(run, size=9, color=RGBColor(0x9A, 0xA8, 0xBC))
        elif name in ("div", "section"):
            walk(doc, BeautifulSoup(str(el), "html.parser"))


def main():
    text = MD_PATH.read_text(encoding="utf-8")
    html = md.markdown(
        text,
        extensions=["tables", "fenced_code", "sane_lists"],
        output_format="html",
    )
    soup = BeautifulSoup(f"<body>{html}</body>", "html.parser")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    walk(doc, soup)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"OK -> {OUT_PATH}")
    print(f"size = {OUT_PATH.stat().st_size} bytes")


if __name__ == "__main__":
    main()

# -*- coding: utf-8 -*-
"""Generate Yang Zhen resume Word doc matching the screenshot layout."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw

ROOT = Path(r"d:\code\AICODES")
SRC_IMG = Path(
    r"C:\Users\yz\.cursor\projects\d-code-AICODES\assets"
    r"\c__Users_yz_AppData_Roaming_Cursor_User_workspaceStorage_"
    r"60154ae4e435d16e10c47968b25cc0db_images______20260830194653_53_2-8adf9e78-2224-48ba-9f89-970efd677f00.jpg"
)
OUT_DIR = ROOT / "简历"
OUT_DOCX = OUT_DIR / "杨珍_工业工程师_简历.docx"
AVATAR_PATH = OUT_DIR / "avatar.png"

BLUE = "2F6FED"
BLUE_RGB = RGBColor(0x2F, 0x6F, 0xED)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_run_font(run, size=10, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def clear_paragraph(p):
    p.clear()


def add_para(cell_or_doc, text="", size=9, bold=False, color=None, align=None, space_after=2, space_before=0):
    if hasattr(cell_or_doc, "paragraphs") and hasattr(cell_or_doc, "add_paragraph"):
        # cell
        p = cell_or_doc.add_paragraph()
    else:
        p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_runs(p, parts):
    """parts: list of (text, size, bold, color)"""
    for text, size, bold, color in parts:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def extract_avatar():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    im = Image.open(SRC_IMG).convert("RGB")
    arr = __import__("numpy").array(im)
    sub = arr[5:200, 10:250]
    is_blue = (sub[:, :, 2] > 150) & (sub[:, :, 2] > sub[:, :, 0] + 40) & (sub[:, :, 0] < 80)
    non = ~is_blue
    ys, xs = __import__("numpy").where(non)
    cx = int((xs.min() + xs.max()) // 2)
    cy = int((ys.min() + ys.max()) // 2)
    r = int(max(xs.max() - xs.min(), ys.max() - ys.min()) // 2 + 4)
    fx, fy = 10 + cx, 5 + cy
    crop = im.crop((fx - r, fy - r, fx + r, fy + r)).resize((256, 256), Image.Resampling.LANCZOS)
    mask = Image.new("L", (256, 256), 0)
    ImageDraw.Draw(mask).ellipse((0, 0, 255, 255), fill=255)
    out = Image.new("RGBA", (256, 256), (0, 0, 0, 0))
    out.paste(crop, (0, 0))
    out.putalpha(mask)
    out.save(AVATAR_PATH)
    return AVATAR_PATH


def build():
    avatar = extract_avatar()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.6)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.6)
    section.right_margin = Cm(0.6)

    # Single full-page 2-col table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(6.4)
    right.width = Cm(13.4)
    set_cell_shading(left, BLUE)
    set_cell_shading(right, "FFFFFF")
    set_cell_margins(left, 80, 80, 100, 90)
    set_cell_margins(right, 60, 60, 120, 80)

    # Clear default empty paragraph in cells carefully later
    # ---- LEFT SIDEBAR ----
    # remove default empty first para content by rewriting
    p0 = left.paragraphs[0]
    clear_paragraph(p0)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p0.add_run()
    run.add_picture(str(avatar), width=Cm(2.6))
    p0.paragraph_format.space_after = Pt(4)

    add_para(left, "杨珍", size=22, bold=True, color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    info_items = [
        ("电话", "13472595846"),
        ("邮箱", "932574166@qq.com"),
        ("性别", "女"),
        ("年龄", "30"),
        ("户籍", "广西壮族自治区"),
        ("现所在地", "上海松江"),
        ("开始工作时间", "2018"),
        ("最高学历", "本科"),
        ("政治面貌", "共青团员"),
        ("民族", "汉族"),
        ("毕业时间", "2018"),
        ("工作年限", "8"),
    ]
    for label, value in info_items:
        p = add_para(left, space_after=1)
        add_runs(
            p,
            [
                (f"• {label}：", 8, False, RGBColor(0xD6, 0xE4, 0xFF)),
                (value, 8, False, RGBColor(255, 255, 255)),
            ],
        )

    add_para(left, "◎ 求职意向", size=11, bold=True, color=RGBColor(255, 255, 255), space_before=10, space_after=4)
    intent = [
        ("意向岗位", "工业工程师(IE)"),
        ("意向城市", "上海"),
        ("期望月薪", "面议"),
        ("求职类型", "社招"),
        ("当前状态", "考虑机会"),
    ]
    for label, value in intent:
        p = add_para(left, space_after=1)
        add_runs(
            p,
            [
                (f"{label}：", 8, False, RGBColor(0xD6, 0xE4, 0xFF)),
                (value, 8, True, RGBColor(255, 255, 255)),
            ],
        )

    add_para(left, "◎ 自我评价", size=11, bold=True, color=RGBColor(255, 255, 255), space_before=10, space_after=4)
    add_para(
        left,
        "1、有7年多的 IE 经验，主导工厂人效改善和管控；厂房规划布局；设备 OEE 分析&产能提升；MES 数字化改善；新设备需求评审和节拍验收等。",
        size=8,
        color=RGBColor(255, 255, 255),
        space_after=4,
    )

    # ---- RIGHT CONTENT ----
    # clear default
    rp0 = right.paragraphs[0]
    clear_paragraph(rp0)

    def section_title(cell, title: str):
        p = add_para(cell, space_before=2, space_after=4)
        add_runs(p, [("●  ", 12, True, BLUE_RGB), (title, 13, True, BLUE_RGB)])
        # underline-ish spacer
        line = add_para(cell, "━" * 36, size=7, color=RGBColor(0xB8, 0xCC, 0xF5), space_after=6)
        return p

    def job_header(cell, period: str, org: str):
        p = add_para(cell, space_after=1, space_before=2)
        # period left, org right via tab-ish spacing
        run1 = p.add_run(period)
        set_run_font(run1, size=9, bold=False, color=GRAY)
        # fill spaces approximately
        run_mid = p.add_run("    ")
        set_run_font(run_mid, size=9)
        run2 = p.add_run(org)
        set_run_font(run2, size=11, bold=True, color=DARK)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Better: use a 2-col nested table for true left/right
        return p

    def header_row(cell, period: str, org: str):
        t = cell.add_table(rows=1, cols=2)
        t.autofit = True
        c0, c1 = t.cell(0, 0), t.cell(0, 1)
        c0.width = Cm(4.2)
        c1.width = Cm(8.6)
        p0 = c0.paragraphs[0]
        clear_paragraph(p0)
        r = p0.add_run(period)
        set_run_font(r, size=9, color=GRAY)
        p1 = c1.paragraphs[0]
        clear_paragraph(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p1.add_run(org)
        set_run_font(r2, size=11, bold=True, color=DARK)

    # Education
    section_title(right, "教育经历")
    header_row(right, "2014.01 - 2018.01", "石河子大学")
    add_para(right, "机械设计制造及其自动化  |  本科", size=9, color=DARK, space_after=8)

    # Work
    section_title(right, "工作经历")
    header_row(right, "2019.05 - 至今", "上海比亚迪有限公司")
    add_para(right, "高级IE工程师", size=9, bold=True, color=BLUE_RGB, space_after=3)
    add_para(right, "内容：", size=9, bold=True, color=DARK, space_after=1)
    byd_content = [
        "工厂人效改善；",
        "厂房布局规划；",
        "线平衡分析；",
        "MES 数字化现场改善；",
        "新设备功能需求评审和节拍验收；",
        "现场5S管理；",
        "SOP 审核及自检基准书编写。",
    ]
    for i, t in enumerate(byd_content, 1):
        add_para(right, f"{i}. {t}", size=9, color=DARK, space_after=0)
    add_para(right, "业绩：", size=9, bold=True, color=DARK, space_before=4, space_after=1)
    byd_perf = [
        "工厂人效改善：制定人员配置标准，工厂万支人员完成20%降幅目标；",
        "外观改善项目：不良率1.49%降至0.6%；",
        "负责A客户新项目布局规划；",
        "主导工厂 OEE 改善和自动化推进；",
        "参与新设备的功能需求评审，以及设备的验收；",
        "新项目导入的成本核算。",
    ]
    for i, t in enumerate(byd_perf, 1):
        add_para(right, f"{i}. {t}", size=9, color=DARK, space_after=0)

    add_para(right, "", size=6, space_before=6, space_after=2)
    header_row(right, "2018.09 - 2019.05", "青岛华旗科技有限公司")
    add_para(right, "机械工程师", size=9, bold=True, color=BLUE_RGB, space_after=3)
    huaqi = [
        "主要负责设备气路控制部分的设计，包括三维模型和二维图纸，以及设备管路的设计。",
        "整个设备气路部分的图纸整理、归档，以及 BOM 表格。",
        "气路件的选型，与厂家的沟通以及部分零件的验收与检测。",
        "主要从事半导体行业非标设备的设计。",
    ]
    for i, t in enumerate(huaqi, 1):
        add_para(right, f"{i}. {t}", size=9, color=DARK, space_after=0)

    # Projects — overview on page 1
    add_para(right, "", size=6, space_before=8, space_after=2)
    section_title(right, "项目经历（概览）")
    overview = [
        ("2023.01 - 至今", "工厂人效改善", "上海工厂人效改善责任人"),
        ("2021.05 - 2023.11", "先进先出MES管理", "项目负责人"),
        ("2020.04 - 2021.01", "单拉产出提升", "项目负责人"),
        ("2019.05 - 2020.12", "外观改善", "项目负责人"),
        ("2019.01 - 2019.03", "12寸高温氧化炉", "机械设计"),
    ]
    for period, name, role in overview:
        header_row(right, period, name)
        add_para(right, role, size=8, color=BLUE_RGB, space_after=3)

    add_para(right, "详见后续「项目经历详情」页。", size=8, color=GRAY, space_before=4)

    # ---- DETAIL PAGES (single column) ----
    def doc_section(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        add_runs(p, [("●  ", 12, True, BLUE_RGB), (title, 14, True, BLUE_RGB)])
        line = doc.add_paragraph("━" * 42)
        line.paragraph_format.space_after = Pt(6)
        for run in line.runs:
            set_run_font(run, size=7, color=RGBColor(0xB8, 0xCC, 0xF5))

    def doc_header(period: str, name: str):
        t = doc.add_table(rows=1, cols=2)
        c0, c1 = t.cell(0, 0), t.cell(0, 1)
        c0.width = Cm(5.5)
        c1.width = Cm(13.5)
        p0 = c0.paragraphs[0]
        clear_paragraph(p0)
        r = p0.add_run(period)
        set_run_font(r, size=10, color=GRAY)
        p1 = c1.paragraphs[0]
        clear_paragraph(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p1.add_run(name)
        set_run_font(r2, size=12, bold=True, color=DARK)

    def doc_line(text: str, size=9, bold=False, color=None, before=0, after=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color or DARK)

    def doc_bullets(items, size=9):
        for i, t in enumerate(items, 1):
            doc_line(f"{i}. {t}", size=size, after=0)

    # Force new page for detailed projects
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)

    doc_section("项目经历详情")

    # 1 工厂人效改善
    doc_header("2023.01 - 至今", "工厂人效改善")
    doc_line("上海工厂人效改善责任人", size=10, bold=True, color=BLUE_RGB, before=2, after=3)
    doc_line("内容：", bold=True, after=2)
    doc_line(
        "项目背景：事业部战略指标下发到工厂，直接设定工厂万支生产人员人效目标。",
        after=1,
    )
    doc_line(
        "项目目标：2023全年人效降幅20%；2024年在2023年最优基础上再降幅20%。",
        after=1,
    )
    doc_line(
        "项目范围：上海工厂直接人员（IH级）；管理人员需做组织瘦身和人员对标。",
        after=1,
    )
    doc_line("主要业绩：2023、2024年平均超额完成人效目标。", after=3)
    doc_line("人效项目主要开展方向：", bold=True, after=2)
    doc_bullets(
        [
            "目标分解：制定考核标准、KPI考核；",
            "组织瘦身：制定人员配置标准，定员定岗；",
            "流程优化：消除浪费，减员增效；",
            "系统监控：开发MES人效系统，实时监控，闭环管理；",
            "人员管控：人员增补标准化、岗位编制管控；",
            "工时管控：制定标准工时库，管控每日出勤工时和库存工时。",
        ]
    )
    doc_line("项目成果：", bold=True, before=4, after=2)
    doc_bullets(
        [
            "通过AGV物流优化、流程优化、数字化改善、自动化改善、精益优化等，累计节省人力390人；",
            "2023年人效目标达成率103%；2024年人效达成连续4个月超过100%；",
            "制定工厂人员配置标准库，以及各机种各拉线标准工时库；",
            "完成人效系统的开发和运用。",
        ]
    )

    # 2 MES FIFO
    doc_line("", before=8)
    doc_header("2021.05 - 2023.11", "先进先出MES管理")
    doc_line("项目负责人", size=10, bold=True, color=BLUE_RGB, before=2, after=3)
    doc_line("1. 项目效果：", bold=True, after=2)
    doc_bullets(
        [
            "老化先进先出管理看板；",
            "各工序在制品库存看板；",
            "各工序防呆；",
            "取消大部分人工统计工作，系统自动生成生产报表；",
            "看板展示每小时单拉产出、不良趋势、设备OEE。",
        ]
    )
    doc_line("2. 项目成员：IE工程师、Java工程师。", before=3, after=2)
    doc_line("3. 项目分配：", bold=True, before=2, after=2)
    doc_bullets(
        [
            "主导：项目负责人；",
            "需求制定：各车间IE工程师；",
            "需求方案：工程师讨论，项目负责人编写具体方案步骤；",
            "需求优先级：组织讨论各需求作用，分析需求间关联与难易度，会上汇报领导。",
        ]
    )
    doc_line("4. 项目进展（现成果）：", bold=True, before=3, after=2)
    doc_bullets(
        [
            "老化先进先出看板已投入使用；",
            "智能转运单完成，每班每区域节省转运工2人、转运文员2人/天；",
            "在制品库存看板完成，每日监控各工序WIP及准时交付完成情况；",
            "生产报表自动生成完成，取消人工报表。",
        ]
    )

    # 3 单拉产出提升
    doc_line("", before=8)
    doc_header("2020.04 - 2021.01", "单拉产出提升")
    doc_line("项目负责人", size=10, bold=True, color=BLUE_RGB, before=2, after=3)
    doc_line("内容：", bold=True, after=2)
    doc_line("1. 项目背景：客户需求增加，产线需做换型及产出提升。", after=2)
    doc_line("项目效果：单拉产出由8500pcs/day提升至9000pcs/day。", after=2)
    doc_line(
        "项目思路：现状调查 → 要因分析 → 措施实施 → 效果确认 → 总结 → 有效措施横向展开。",
        after=2,
    )
    doc_line("项目主要改善点：", bold=True, after=2)
    doc_bullets(
        [
            "分析设备OEE，制定统一数据统计模板，定期分析并讨论TOP3问题；",
            "分析人员作业内容，测定作业频次与节拍，从人员侧消除影响产出的因素（如降低换料频次）；",
            "分析瓶颈工序，测定工序节拍，从设备动作与工艺标准侧消除影响因素（如缩短压芯时长）。",
        ]
    )
    doc_line("业绩：", bold=True, before=3, after=2)
    doc_bullets(
        [
            "项目收益：9万元/月；",
            "项目结果：8500pcs/day → 9000pcs/day。",
        ]
    )

    # 4 外观改善
    doc_line("", before=8)
    doc_header("2019.05 - 2020.12", "外观改善")
    doc_line("项目负责人", size=10, bold=True, color=BLUE_RGB, before=2, after=3)
    doc_line("内容：", bold=True, after=2)
    doc_bullets(
        [
            "改善效果：外观不良率由1.49%降至0.6%，并维持5个月以上；",
            "项目收益：每月节省约12万元；",
            "改善思路：现状分析（近阶段生产数据分析，找出TOP5外观不良类型）→ 要因确认（按工艺流程逐段分析主要因素，找改善切入点）→ 分析改善（观察设备动作并结合工作原理，分析具体动作与改善方法）→ 措施改善与巩固（设备侧：增加小治具机构、纳入PM要求；作业侧：通过SOP/自检基准书/一点课等文件与培训，规范员工操作）；",
            "改善实例：内部异物——增加吹气装置并更新点检表；表面白点——增加刮刀挡板并提高刮刀更换频次；表面凹坑——制定自检基准书与换型管理规定。",
        ]
    )
    doc_line("业绩：", bold=True, before=3, after=2)
    doc_bullets(
        [
            "项目收益：12万元/月；",
            "项目结果：不良率1.49% → 0.6%。",
        ]
    )

    # 5 氧化炉
    doc_line("", before=8)
    doc_header("2019.01 - 2019.03", "12寸高温氧化炉")
    doc_line("机械设计", size=10, bold=True, color=BLUE_RGB, before=2, after=3)
    doc_line("内容：", bold=True, after=2)
    doc_line(
        "为中科院苏州纳米所提供设备；与12寸高温退火炉同步开展设计。",
        after=2,
    )
    doc_line("业绩：按时完成任务。", after=6)

    # Certificates & skills
    doc_section("荣誉证书")
    doc_line("大学英语四级、PMP项目管理认证、六西格玛培训结业证书。", after=8)

    doc_section("相关技能")
    doc_line("3D建模、CAD、项目管理、六西格玛。", after=4)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(OUT_DOCX.stat().st_size)


if __name__ == "__main__":
    build()
